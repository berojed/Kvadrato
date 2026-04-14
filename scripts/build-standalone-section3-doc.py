#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_DOCX = Path("/Users/bernard/Desktop/Kvadrato/Dokumentacija/Section_3_Prototip_UIUX_standalone.docx")
OUT_MD = Path("/Users/bernard/Desktop/Kvadrato/Dokumentacija/Section_3_Prototip_UIUX_standalone.md")
OUT_MAPPING = Path("/Users/bernard/Desktop/Kvadrato/Dokumentacija/Section_3_Prototip_UIUX_mapping.txt")

PROTOTYPE_DIR = Path("/Users/bernard/Desktop/Kvadrato/Dokumentacija/slike/prototype")
REAL_DIR = Path("/Users/bernard/Desktop/Kvadrato/Dokumentacija/slike/section3-real-2026-04-07")


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_figure(
    doc: Document,
    image_path: Path,
    figure_no: int,
    caption_text: str,
    explanation_text: str,
) -> None:
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))

    p_cap = doc.add_paragraph(f"Slika {figure_no}: {caption_text}")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(doc, explanation_text)


def main() -> None:
    prototype_figures = [
        {
            "path": PROTOTYPE_DIR / "00_flow_selector.png",
            "caption": "Početni prototipski odabir korisničkog toka (kupac/prodavatelj)",
            "explanation": "Početni ekran prototipa formalizira grananje korisničkih putanja prema ulozi, što je omogućilo ranu validaciju informacijske arhitekture sustava. Takav pristup je važan jer smanjuje rizik pogrešne navigacije u kasnijoj implementaciji i olakšava dosljedno oblikovanje funkcionalnosti po korisničkim profilima.",
        },
        {
            "path": PROTOTYPE_DIR / "01_buyer_home.png",
            "caption": "Prototipski prikaz početne stranice kupca",
            "explanation": "Prikaz početne stranice definira hijerarhiju sadržaja: istaknuta vrijednosna poruka, pretraga i pregled ključnih oglasa. Time je u ranoj fazi uspostavljen model početne orijentacije korisnika, što je kritično za vrijeme do prve relevantne interakcije.",
        },
        {
            "path": PROTOTYPE_DIR / "02_buyer_listings.png",
            "caption": "Prototipski pregled oglasa s filtracijskim panelom",
            "explanation": "Na ovoj slici prikazana je kombinacija liste nekretnina i bočnog filtriranja, čime se modelira glavni istraživački tok korisnika. Integracija filtera uz rezultate omogućila je validaciju kognitivnog opterećenja i učinkovitosti sužavanja skupa rezultata.",
        },
        {
            "path": PROTOTYPE_DIR / "03_buyer_detail.png",
            "caption": "Prototipski detaljni prikaz pojedine nekretnine",
            "explanation": "Detaljni prikaz strukturira sve informacije potrebne za evaluaciju oglasa: medijski sadržaj, tehničke podatke, lokaciju i kontaktne akcije. Ovakva organizacija podržava donošenje informirane odluke bez fragmentiranja korisničke pažnje kroz više nepovezanih zaslona.",
        },
        {
            "path": PROTOTYPE_DIR / "04_buyer_viewer3d.png",
            "caption": "Prototipski prikaz interaktivnog 3D pregleda nekretnine",
            "explanation": "3D prikaz je u prototipu definiran kao diferencirajuća funkcionalnost koja proširuje klasičan 2D pregled fotografija. Uključivanjem ove komponente u ranoj fazi testirana je izvedivost naprednijeg korisničkog iskustva i njegova pozicija unutar osnovnog toka pregledavanja.",
        },
        {
            "path": PROTOTYPE_DIR / "08_seller_dashboard.png",
            "caption": "Prototipski prikaz kontrolne ploče prodavatelja",
            "explanation": "Kontrolna ploča prodavatelja prikazuje metrika-orijentiran pristup upravljanju oglasima i aktivnostima korisnika. Time je u prototipu uspostavljen operativni model rada za poslovnog korisnika, što je kasnije preneseno u implementaciju administrativnog sučelja.",
        },
    ]

    real_figures = [
        {
            "path": REAL_DIR / "kvadrato_real_homepage_balanced.png",
            "caption": "Početna stranica implementirane aplikacije Kvadrato",
            "explanation": "Početni zaslon u produkcijskoj verziji zadržava jasnu ulaznu točku i istaknuti pretraživački mehanizam. Takav raspored poboljšava početnu upotrebljivost jer korisniku odmah komunicira glavnu svrhu sustava i omogućuje brz ulazak u proces pronalaska nekretnine.",
        },
        {
            "path": REAL_DIR / "kvadrato_real_authentication_login.png",
            "caption": "Autentikacijski zaslon s odabirom korisničke uloge",
            "explanation": "Autentikacijski tok eksplicitno razdvaja uloge kupca i prodavatelja prije pristupa računima. Ovakvo rješenje je funkcionalno važno jer uvjetuje prikaz kontekstualno relevantnih opcija i smanjuje broj pogrešnih radnji nakon prijave.",
        },
        {
            "path": REAL_DIR / "kvadrato_real_listing_overview.png",
            "caption": "Implementirani pregled oglasa s realnim podacima nekretnina",
            "explanation": "Prikaz oglasa demonstrira rad sustava nad produkcijski vjerodostojnim podacima, uključujući naslov, lokaciju, cijenu i površinu nekretnine. Korisnik kroz jedinstveni ekran može usporediti više kandidata i brzo odabrati one koji odgovaraju njegovim kriterijima.",
        },
        {
            "path": REAL_DIR / "kvadrato_real_filters_search.png",
            "caption": "Pretraga i filtriranje ponude nekretnina u implementaciji",
            "explanation": "Na slici je vidljiva kombinacija tekstualnog pretraživanja i atributnih filtera (tip oglasa, cijena, broj soba), što omogućuje precizno sužavanje rezultata. Ovakav model interakcije smanjuje vrijeme pronalaska relevantnog oglasa i povećava učinkovitost rada korisnika.",
        },
        {
            "path": REAL_DIR / "kvadrato_real_property_detail.png",
            "caption": "Detaljni prikaz nekretnine s realnim sadržajem oglasa",
            "explanation": "Detaljni ekran prikazuje nekretninu „Moderan dvosoban stan s balkonom u centru Zagreba“ uz medijski sadržaj i akcije kontakta. Integracija informacija i akcijskih elemenata na istom zaslonu podržava prijelaz iz pasivnog pregleda u aktivnu korisničku namjeru (upit ili razgledavanje).",
        },
        {
            "path": REAL_DIR / "kvadrato_real_settings_preferences.png",
            "caption": "Postavke sučelja (jezik, tema i tipografija) u produkcijskoj aplikaciji",
            "explanation": "Dodatni važan zaslon prikazuje personalizaciju vizualnog sučelja kroz jezik, temu i font. Time se povećava prilagodljivost sustava različitim korisničkim preferencijama i osigurava konzistentnije iskustvo kroz različite sesije korištenja.",
        },
    ]

    for entry in [*prototype_figures, *real_figures]:
        if not entry["path"].exists():
            raise FileNotFoundError(f"Missing screenshot: {entry['path']}")

    doc = Document()
    set_default_style(doc)

    add_heading(doc, "3. Prototip i UI/UX implementacija aplikacije", 1)
    add_body(
        doc,
        "Ovo poglavlje sažeto prikazuje prijelaz od dizajnerskog prototipa do završne implementacije aplikacije Kvadrato. Naglasak je na funkcionalnoj i uporabnoj vrijednosti sučelja te na povezanosti između početnih UX odluka i produkcijskog rješenja.",
    )

    add_heading(doc, "3.1 Prototip aplikacije", 2)
    add_body(
        doc,
        "Prototip razvijen u datoteci Kvadrato_Prototip.sketch korišten je za ranu validaciju rasporeda elemenata, navigacijskog toka i ključnih interakcija. U ovoj fazi fokus nije bio na tehničkoj izvedbi nego na strukturi korisničkog iskustva i logici kretanja kroz sustav.",
    )

    figure_no = 1
    mapping_lines = [
        "Standalone Section 3 mapping",
        f"Output DOCX: {OUT_DOCX}",
        "",
        "PROTOTIP (6):",
    ]
    for fig in prototype_figures:
        add_figure(doc, fig["path"], figure_no, fig["caption"], fig["explanation"])
        mapping_lines.append(f"Slika {figure_no} -> {fig['path'].name}")
        figure_no += 1

    add_heading(doc, "3.2 UI/UX gotovog produkta", 2)
    add_body(
        doc,
        "Implementirana aplikacija razvijena je u React okruženju i koristi produkcijski vjerodostojne podatke oglasa. U nastavku su prikazani ključni zasloni koji potvrđuju ostvarenje osnovnih korisničkih tokova: početna orijentacija, autentikacija, pretraživanje, filtriranje i detaljna evaluacija nekretnine.",
    )

    mapping_lines.extend(["", "REAL APP (6):"])
    for fig in real_figures:
        add_figure(doc, fig["path"], figure_no, fig["caption"], fig["explanation"])
        mapping_lines.append(f"Slika {figure_no} -> {fig['path'].name}")
        figure_no += 1

    add_heading(doc, "3.3 Usporedba prototipa i implementacije", 2)
    add_body(
        doc,
        "Temeljna struktura sučelja ostala je dosljedna prototipu: početni zaslon kao ulazna točka, centralni pregled oglasa i detaljni prikaz nekretnine kao glavna odluka korisničkog toka. Time je potvrđena ispravnost ranih UX pretpostavki postavljenih u Sketch fazi.",
    )
    add_body(
        doc,
        "Ključna razlika odnosi se na razinu operativnosti. Prototip je demonstrirao logiku i raspored, dok implementacija uvodi stvarne podatkovne veze, autentikaciju po ulogama te aktivno filtriranje rezultata. Zbog tehničkih ograničenja produkcijskog sustava uvedene su i dodatne kontrole stanja, validacije i prilagodbe performansi.",
    )
    add_body(
        doc,
        "Iz perspektive uporabljivosti, finalna verzija unaprijeđena je kroz bržu orijentaciju, preciznije pronalaženje relevantnih oglasa i jasnije akcijske točke na detalju nekretnine. Time je prototipska koncepcija transformirana u stabilno i funkcionalno rješenje prikladno za realne korisničke scenarije.",
    )

    doc.save(str(OUT_DOCX))

    md_lines = [
        "# 3. Prototip i UI/UX implementacija aplikacije",
        "",
        "Ovo poglavlje sažeto prikazuje prijelaz od dizajnerskog prototipa do završne implementacije aplikacije Kvadrato. Naglasak je na funkcionalnoj i uporabnoj vrijednosti sučelja te na povezanosti između početnih UX odluka i produkcijskog rješenja.",
        "",
        "## 3.1 Prototip aplikacije",
        "",
        "Prototip razvijen u datoteci `Kvadrato_Prototip.sketch` korišten je za ranu validaciju rasporeda elemenata, navigacijskog toka i ključnih interakcija. U ovoj fazi fokus nije bio na tehničkoj izvedbi nego na strukturi korisničkog iskustva i logici kretanja kroz sustav.",
        "",
    ]

    def md_figure(no: int, fig: dict) -> list[str]:
        return [
            f"![Slika {no}]({fig['path']})",
            f"*Slika {no}: {fig['caption']}*",
            "",
            fig["explanation"],
            "",
        ]

    no = 1
    for fig in prototype_figures:
        md_lines.extend(md_figure(no, fig))
        no += 1

    md_lines.extend(
        [
            "## 3.2 UI/UX gotovog produkta",
            "",
            "Implementirana aplikacija razvijena je u React okruženju i koristi produkcijski vjerodostojne podatke oglasa. U nastavku su prikazani ključni zasloni koji potvrđuju ostvarenje osnovnih korisničkih tokova: početna orijentacija, autentikacija, pretraživanje, filtriranje i detaljna evaluacija nekretnine.",
            "",
        ]
    )

    for fig in real_figures:
        md_lines.extend(md_figure(no, fig))
        no += 1

    md_lines.extend(
        [
            "## 3.3 Usporedba prototipa i implementacije",
            "",
            "Temeljna struktura sučelja ostala je dosljedna prototipu: početni zaslon kao ulazna točka, centralni pregled oglasa i detaljni prikaz nekretnine kao glavna odluka korisničkog toka. Time je potvrđena ispravnost ranih UX pretpostavki postavljenih u Sketch fazi.",
            "",
            "Ključna razlika odnosi se na razinu operativnosti. Prototip je demonstrirao logiku i raspored, dok implementacija uvodi stvarne podatkovne veze, autentikaciju po ulogama te aktivno filtriranje rezultata. Zbog tehničkih ograničenja produkcijskog sustava uvedene su i dodatne kontrole stanja, validacije i prilagodbe performansi.",
            "",
            "Iz perspektive uporabljivosti, finalna verzija unaprijeđena je kroz bržu orijentaciju, preciznije pronalaženje relevantnih oglasa i jasnije akcijske točke na detalju nekretnine. Time je prototipska koncepcija transformirana u stabilno i funkcionalno rješenje prikladno za realne korisničke scenarije.",
            "",
        ]
    )

    OUT_MD.write_text("\n".join(md_lines), encoding="utf-8")
    OUT_MAPPING.write_text("\n".join(mapping_lines), encoding="utf-8")

    print(f"Created DOCX: {OUT_DOCX}")
    print(f"Created MD: {OUT_MD}")
    print(f"Created mapping: {OUT_MAPPING}")


if __name__ == "__main__":
    main()
